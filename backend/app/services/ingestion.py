import os, tempfile, logging, uuid, re
from types import SimpleNamespace as NS
from typing import List, Dict, Any, Optional, Tuple
import boto3
from botocore.config import Config as BotoConfig
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue, FilterSelector, PayloadSchemaType
from openai import OpenAI
from app.core.config import settings

# Word document support
def _check_docx_available():
    try:
        from docx import Document
        return True
    except ImportError:
        return False

# OCR imports - check dynamically to avoid caching issues
def _check_ocr_available():
    try:
        import pytesseract
        from pdf2image import convert_from_path
        from PIL import Image
        print("✅ OCR libraries successfully imported in _check_ocr_available()")
        return True
    except ImportError as e:
        print(f"❌ OCR import failed: {e}")
        return False
    except Exception as e:
        print(f"❌ OCR check unexpected error: {e}")
        return False

log = logging.getLogger("ingestion")
log.setLevel(logging.INFO)

STATUS: Dict[str, Dict[str, Any]] = {}
EMBED_MODEL = "text-embedding-3-small"
PRECEDENCE = {"AAOS":100, "RCT":100, "CLINICAL_GUIDELINE":100, "HOSPITAL_POLICY":90, "PEER_REVIEW":80, "DOCTOR_PROTOCOL":95, "OTHER":50}

# Mapping from doctor slug (as used in collection names) to display name.
# For protocol documents, the author is the doctor themselves — not whatever
# is embedded in the PDF metadata (which is often the typist or blank).
DOCTOR_SLUG_TO_NAME: Dict[str, str] = {
    "joshua_dines": "Dr. Joshua Dines",
    "asheesh_bedi": "Dr. Asheesh Bedi",
    "ayoosh_pareek": "Dr. Ayoosh Pareek",
    "sheeraz_qureshi": "Dr. Sheeraz Qureshi",
    "khalid_alkhelaifi": "Dr. Khalid Alkhelaifi",
    "william_long": "Dr. William Long",
    "jorge_chahla": "Dr. Jorge Chahla",
    "steven_defroda": "Dr. Steven DeFroda",
}

# Source types that represent research / external evidence.
# These keep their extracted PDF metadata author (the paper's actual authors).
_RESEARCH_SOURCE_TYPES = frozenset({"AAOS", "RCT", "CLINICAL_GUIDELINE", "PEER_REVIEW"})


def _resolve_protocol_author(collection_name: str, source_type: str) -> Optional[str]:
    """Return the doctor's display name when the document is a doctor protocol.

    For doctor-specific collections (not ``dr_general_*``) where the source
    type is *not* a research article, the author should be the doctor who
    owns the collection — regardless of what the PDF metadata says.

    Research articles (RCTs, AAOS guidelines, etc.) keep their extracted
    author metadata so that the citation reads correctly.
    """
    if source_type.upper() in _RESEARCH_SOURCE_TYPES:
        return None
    if not collection_name.startswith("dr_") or collection_name.startswith("dr_general_"):
        return None
    for slug, name in DOCTOR_SLUG_TO_NAME.items():
        if collection_name.startswith(f"dr_{slug}_"):
            return name
    return None

def _s3():
    return boto3.client("s3", region_name=settings.aws_region, config=BotoConfig(retries={"max_attempts": 3}))

def _qdrant():
    kw = dict(url=settings.qdrant_url, timeout=90)
    if settings.qdrant_api_key: kw["api_key"] = settings.qdrant_api_key
    return QdrantClient(**kw)

def _openai(): 
    return OpenAI(api_key=settings.openai_api_key)

def _ensure_collection(cli: QdrantClient, collection_name: str):
    try:
        cli.get_collection(collection_name)
    except Exception:
        cli.recreate_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=1536, distance=Distance.COSINE)
        )
    # Ensure a keyword index on document_id so we can filter/delete by it.
    # create_payload_index is idempotent — no-ops if it already exists.
    try:
        cli.create_payload_index(
            collection_name=collection_name,
            field_name="document_id",
            field_schema=PayloadSchemaType.KEYWORD,
        )
    except Exception:
        pass  # Index already exists or collection just created with it

def _download(bucket: str, key: str) -> Tuple[str, Optional[str]]:
    """Download file from S3 and return (local_path, original_filename)."""
    head_resp = _s3().head_object(Bucket=bucket, Key=key)
    # Get original filename from S3 metadata
    original_filename = head_resp.get('Metadata', {}).get('original-filename')

    fd, path = tempfile.mkstemp(prefix="doc-", suffix=os.path.splitext(key)[1])
    os.close(fd)
    with open(path, "wb") as f:
        _s3().download_fileobj(bucket, key, f)
    return path, original_filename


def _filename_to_title(filename: str) -> str:
    """Convert a filename to a readable title.

    Examples:
        "Rotator_Cuff_Repair_Protocol.pdf" -> "Rotator Cuff Repair Protocol"
        "UCL-reconstruction-2023.docx" -> "UCL Reconstruction 2023"
    """
    if not filename:
        return "Unknown Document"

    # Remove file extension
    name = os.path.splitext(filename)[0]

    # Replace underscores and hyphens with spaces
    name = name.replace('_', ' ').replace('-', ' ')

    # Clean up multiple spaces
    name = ' '.join(name.split())

    # Title case if all lowercase or all uppercase
    if name.islower() or name.isupper():
        name = name.title()

    return name.strip() if name.strip() else "Unknown Document"

def _ocr_pdf(path: str) -> List[NS]:
    """Extract text from scanned PDF using OCR."""
    if not _check_ocr_available():
        raise RuntimeError("OCR libraries not installed. Run: pip install pytesseract pdf2image pillow")
    
    # Import here to avoid issues if not installed
    import pytesseract
    from pdf2image import convert_from_path
    
    log.info("Running OCR on scanned PDF...")
    elements = []
    
    try:
        # Convert PDF to images
        images = convert_from_path(path, dpi=300)
        log.info(f"Processing {len(images)} pages with OCR...")
        
        for page_num, image in enumerate(images, start=1):
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image)
            text = text.strip()
            
            if text:
                elements.append(NS(
                    text=text,
                    metadata=NS(page_number=page_num, category="ocr")
                ))
                log.info(f"OCR extracted {len(text)} chars from page {page_num}")
            else:
                log.warning(f"No text found on page {page_num}")
        
        return elements
    except Exception as e:
        log.error(f"OCR failed: {e}")
        raise RuntimeError(f"OCR processing failed: {str(e)}")

def _parse_pdf_fast(path: str) -> List[NS]:
    """Try normal extraction first, fall back to OCR if no text found."""
    r = PdfReader(path)
    out = []
    total_text_length = 0
    
    # Try extracting text normally
    for i, p in enumerate(r.pages):
        t = (p.extract_text() or "").strip()
        if t:
            out.append(NS(text=t, metadata=NS(page_number=i+1, category=None)))
            total_text_length += len(t)
    
    # If very little text extracted (likely scanned), use OCR
    if total_text_length < 100 and len(r.pages) > 0:
        print(f"⚠️  Only {total_text_length} chars extracted from {len(r.pages)} pages - PDF appears to be scanned")
        ocr_available = _check_ocr_available()
        print(f"🔍 OCR check result: {ocr_available}")
        
        if ocr_available:
            print("🚀 Starting OCR extraction...")
            result = _ocr_pdf(path)
            print(f"✅ OCR extracted {len(result)} elements")
            return result
        else:
            print("❌ OCR not available")
            raise RuntimeError("PDF appears to be scanned but OCR is not available. Install: pip install pytesseract pdf2image pillow")
    
    return out

def _parse_docx(path: str) -> List[NS]:
    """Parse a Word document (.docx) and extract text.

    Concatenates all paragraph text into a single string so that the
    downstream ``_split`` function can create coherent chunks that
    preserve section context (e.g., a medication name stays with its
    dosage and timing instructions).

    Without concatenation, each paragraph/bullet becomes its own tiny
    chunk with a separate embedding, causing retrieval to miss related
    details (e.g., returning "Ibuprofen:" but not its dosage).

    If no text is found via paragraphs/tables, falls back to extracting
    all ``<w:t>`` text runs from the document XML.
    """
    if not _check_docx_available():
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    from docx import Document

    doc = Document(path)
    elements = []

    # Collect all paragraph text into a single string.
    # python-docx treats each line/bullet as a separate paragraph;
    # concatenating them lets _split() create coherent chunks that
    # keep related content (e.g., medication name + dosage) together,
    # mirroring how the PDF parser returns all text per page.
    para_texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            para_texts.append(text)

    if para_texts:
        elements.append(NS(
            text="\n".join(para_texts),
            metadata=NS(page_number=1, category=None)
        ))

    # Also extract text from tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))
        if table_text:
            elements.append(NS(
                text="\n".join(table_text),
                metadata=NS(page_number=1, category="table")
            ))

    # Fallback: if no text was found via paragraphs/tables, extract all <w:t>
    # text runs from the raw XML.  This catches content in text boxes, shapes,
    # content controls, headers, footers, and other structures.
    if not elements:
        log.warning("No text from paragraphs/tables in %s – trying raw XML extraction", path)
        try:
            import zipfile
            from xml.etree import ElementTree as ET

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            text_parts: List[str] = []

            with zipfile.ZipFile(path) as zf:
                # Collect XML parts that may contain text
                xml_targets = [
                    n for n in zf.namelist()
                    if n.startswith("word/") and n.endswith(".xml")
                ]
                for target in xml_targets:
                    tree = ET.parse(zf.open(target))
                    for t_el in tree.iter(f"{{{ns['w']}}}t"):
                        if t_el.text:
                            text_parts.append(t_el.text)

            full_text = " ".join(text_parts).strip()
            if full_text:
                elements.append(NS(
                    text=full_text,
                    metadata=NS(page_number=1, category="xml_fallback"),
                ))
                log.info("XML fallback extracted %d chars from %s", len(full_text), path)
        except Exception as exc:
            log.error("XML fallback extraction failed for %s: %s", path, exc)

    return elements

def _extract_docx_metadata(path: str) -> Tuple[Optional[str], Optional[str], Optional[int]]:
    """Extract metadata from Word document."""
    try:
        if not _check_docx_available():
            return (None, None, None)

        from docx import Document
        doc = Document(path)

        title = None
        author = None
        year = None

        # Get metadata from document properties
        core_props = doc.core_properties
        if core_props:
            if core_props.title:
                title = core_props.title
            if core_props.author:
                author = core_props.author
            if core_props.created:
                year = core_props.created.year
            if not year and core_props.modified:
                year = core_props.modified.year

        # Fallback: try to extract from first paragraph
        if not title and doc.paragraphs:
            for para in doc.paragraphs[:3]:
                text = para.text.strip()
                if text and len(text) < 200:
                    title = text
                    break

        return (title, author, year)
    except Exception as e:
        log.error(f"Failed to extract docx metadata: {e}")
        return (None, None, None)

def _extract_metadata(path: str) -> Tuple[Optional[str], Optional[str], Optional[int]]:
    """
    Extract meaningful metadata from PDF.
    Returns: (title, author, publication_year)
    """
    try:
        r = PdfReader(path)
        metadata = r.metadata

        # Initialize with None
        title = None
        author = None
        year = None

        # Try to get metadata from PDF properties
        if metadata:
            # Get title from metadata
            if hasattr(metadata, 'title') and metadata.title:
                title = metadata.title.strip()

            # Get author from metadata
            if hasattr(metadata, 'author') and metadata.author:
                author = metadata.author.strip()

            # Get year from creation/modification date
            if hasattr(metadata, 'creation_date') and metadata.creation_date:
                try:
                    year = metadata.creation_date.year
                except:
                    pass

            # If no creation date, try modification date
            if not year and hasattr(metadata, 'modification_date') and metadata.modification_date:
                try:
                    year = metadata.modification_date.year
                except:
                    pass

        # If we still don't have title or author, try to extract from first page
        if (not title or not author or not year) and len(r.pages) > 0:
            first_page_text = r.pages[0].extract_text() or ""

            # Try to extract title from first page (usually first few lines)
            if not title:
                lines = [l.strip() for l in first_page_text.split('\n') if l.strip()]
                if lines:
                    # Take first non-empty line as potential title
                    potential_title = lines[0]
                    # Only use if it looks like a title (not too long, not a URL, etc.)
                    if len(potential_title) < 200 and not potential_title.startswith('http'):
                        title = potential_title

            # Try to extract author patterns (e.g., "Author Name, MD" or "By Author Name")
            if not author:
                # Look for common author patterns
                author_patterns = [
                    r'(?:By|Author[s]?:|Written by)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
                    r'([A-Z][a-z]+\s+[A-Z][a-z]+)(?:,?\s+(?:MD|PhD|M\.D\.|Ph\.D\.))',
                ]
                for pattern in author_patterns:
                    match = re.search(pattern, first_page_text)
                    if match:
                        author = match.group(1).strip()
                        break

            # Try to extract year (4 digits, typically between 1900-2099)
            if not year:
                year_pattern = r'\b(19\d{2}|20\d{2})\b'
                years = re.findall(year_pattern, first_page_text)
                if years:
                    # Take the first year found
                    year = int(years[0])

        log.info(f"Extracted metadata - Title: {title}, Author: {author}, Year: {year}")
        return (title, author, year)

    except Exception as e:
        log.error(f"Failed to extract metadata: {e}")
        return (None, None, None)

def _split(text: str, max_chars=1800, overlap=200) -> List[str]:
    """Split text into chunks, preferring section/paragraph boundaries over mid-sentence splits.

    For structured protocol documents (e.g., post-op protocols organized by time period),
    this preserves section headers with their content so the LLM can correctly associate
    instructions with the right time period.
    """
    if not text: return []
    chunks = []
    i = 0
    L = len(text)
    while i < L:
        j = min(L, i+max_chars)
        if j >= L:
            piece = text[i:j].strip()
            if piece:
                chunks.append(piece)
            break

        # Priority 1: split at a section header boundary (newline before a header-like line).
        # This keeps headers attached to the content that follows them.
        header_cut = -1
        search_start = i + int(max_chars * 0.4)
        for m in re.finditer(
            r'\n(?=(?:Post[- ]?Op|Phase|Week|Day|Month|Stage|Goal|Precaution|Weight|ROM|Range|Brace|Exercise|Rehab|Return|Activity|Restrict)\b)',
            text[search_start:j],
            re.IGNORECASE,
        ):
            header_cut = search_start + m.start()

        # Priority 2: split at a double newline (paragraph boundary)
        if header_cut == -1:
            para_cut = text.rfind("\n\n", search_start, j)
            if para_cut != -1:
                header_cut = para_cut

        # Priority 3: fall back to sentence boundary
        if header_cut != -1:
            cut = header_cut
        else:
            cut = text.rfind(". ", i, j)
            cut = j if cut == -1 or cut < i+max_chars*0.6 else cut+1

        piece = text[i:cut].strip()
        if piece:
            chunks.append(piece)
        i = max(cut-overlap, 0) if cut != j else j
    return chunks

_SECTION_HEADER_RE = re.compile(
    r'^((?:Post[- ]?Op(?:erative)?|Phase|Week|Day|Month|Stage)\s*[^\n]{0,60})',
    re.IGNORECASE | re.MULTILINE,
)

def _detect_section(text: str) -> Optional[str]:
    """Try to detect a section header at the start of the chunk text."""
    m = _SECTION_HEADER_RE.search(text[:200])
    return m.group(1).strip() if m else None

def _to_chunks(elems: List[NS]) -> List[Dict[str, Any]]:
    out = []
    for el in elems:
        txt = (getattr(el,"text","") or "").strip()
        if not txt: continue
        page = getattr(getattr(el,"metadata",NS()),"page_number",None)
        for sub in _split(txt):
            section = _detect_section(sub)
            out.append({"text":sub,"page":page,"section":section})
    return out

def _embed(texts: List[str]) -> List[List[float]]:
    if not texts: return []
    cli = _openai()
    out = []
    B = 128
    for i in range(0, len(texts), B):
        resp = cli.embeddings.create(model=EMBED_MODEL, input=texts[i:i+B])
        out.extend([d.embedding for d in resp.data])
    return out

def process_document(s3_key: str, source_type: str = "OTHER", org_id: str = "demo"):
    bucket = settings.s3_bucket
    doc_id = s3_key

    # Extract collection name from org_id (which is dr_name_protocol format)
    collection_name = org_id if org_id.startswith("dr_") else settings.collection

    STATUS[doc_id] = {"state":"processing"}
    log.info("INGEST start %s", s3_key)
    try:
        path, original_filename = _download(bucket, s3_key)
        log.info("Downloaded %s (original filename: %s)", path, original_filename)

        # Determine file type and parse accordingly
        file_ext = os.path.splitext(path)[1].lower()

        if file_ext in ['.docx', '.doc']:
            # Word document
            doc_title, doc_author, doc_year = _extract_docx_metadata(path)
            elems = _parse_docx(path)
        else:
            # PDF (default)
            doc_title, doc_author, doc_year = _extract_metadata(path)
            elems = _parse_pdf_fast(path)

        # Fallback to original filename (from S3 metadata) if no title extracted
        if not doc_title:
            if original_filename:
                doc_title = _filename_to_title(original_filename)
            else:
                # Last resort: use S3 key basename
                doc_title = _filename_to_title(os.path.basename(s3_key))

        # For doctor protocol documents, the author is the doctor — not
        # whatever was embedded in the PDF (often a typist or blank).
        # Research articles keep their extracted author metadata.
        protocol_author = _resolve_protocol_author(collection_name, source_type)
        if protocol_author:
            log.info("Overriding author for doctor protocol: %s -> %s", doc_author, protocol_author)
            doc_author = protocol_author

        if not elems:
            raise RuntimeError("No text extracted. Document may be empty or parsing failed.")

        chunks = _to_chunks(elems)
        log.info("Chunks pre-embed=%d", len(chunks))

        vecs = _embed([c["text"] for c in chunks])
        if len(vecs) != len(chunks):
            raise RuntimeError("Embedding mismatch")

        cli = _qdrant()
        _ensure_collection(cli, collection_name)

        # Remove any existing chunks for this document so re-ingestion
        # doesn't create duplicates (e.g., old fragmented chunks alongside
        # new properly-chunked ones).
        try:
            cli.delete(
                collection_name=collection_name,
                points_selector=FilterSelector(
                    filter=Filter(
                        must=[FieldCondition(key="document_id", match=MatchValue(value=doc_id))]
                    )
                ),
            )
            log.info("Cleared existing chunks for %s in %s", doc_id, collection_name)
        except Exception as exc:
            log.warning("Could not clear old chunks for %s: %s", doc_id, exc)

        rank = PRECEDENCE.get(source_type.upper(), PRECEDENCE["OTHER"])

        points = []
        for i, (c, v) in enumerate(zip(chunks, vecs)):
            # Use UUID for point ID instead of string path (Qdrant requires UUID or int)
            pid = str(uuid.uuid4())
            payload = {
                "document_id": doc_id,
                "org_id": org_id,
                "source_type": source_type,
                "precedence": rank,
                "page": c.get("page"),
                "section": c.get("section"),
                "text": c["text"],
                "title": doc_title,
                "author": doc_author,
                "publication_year": doc_year,
                "chunk_index": i,  # Add chunk index to payload for reference
                "original_filename": original_filename,  # Preserve original filename for reference
            }
            points.append(PointStruct(id=pid, vector=v, payload=payload))
        
        print(f"📤 Attempting to upsert {len(points)} points to {collection_name}")
        try:
            cli.upsert(collection_name=collection_name, points=points)
            print(f"✅ Successfully upserted {len(points)} points")
        except Exception as e:
            print(f"❌ Upsert failed: {type(e).__name__}: {e}")
            raise
        STATUS[doc_id] = {"state":"done","chunks":len(chunks),"vectors":len(vecs)}
        log.info("INGEST done %s chunks=%d", s3_key, len(chunks))
        
    except Exception as e:
        STATUS[doc_id] = {"state":"error","error":str(e)}
        log.exception("INGEST failed %s: %s", s3_key, e)
    finally:
        # Clean up temp file
        try:
            if 'path' in locals():
                os.unlink(path)
        except:
            pass
